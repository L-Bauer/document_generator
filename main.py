from docx import Document
import os


def main():
    position = "Engineer"
    company = "Example Name"
    template_file_path = "Cover Letter.docx"
    output_file_path = f"Cover Letter - {company}.docx"

    variables = {
        "${COMPANY}": company,
        "${POSITION}": position,
    }

    template_document = Document(template_file_path)

    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(
                            paragraph, variable_key, variable_value
                        )

    template_document.save(output_file_path)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


if __name__ == "__main__":
    main()
